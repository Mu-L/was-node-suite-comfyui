"""A document written as a ``.docx`` file through python-docx.

Every block and run is mapped onto the python-docx object model. Metadata goes to the core
properties, and custom properties through :mod:`.ooxml`.
"""

from __future__ import annotations

from dataclasses import replace
from datetime import datetime, timezone
from io import BytesIO
from typing import TYPE_CHECKING, Any

from ... import deps, log
from .. import metadata as metadata_module
from ..summary import keywords_text
from . import FEATURE, Page, blocks, ooxml

if TYPE_CHECKING:  # pragma: no cover - imported for the annotation alone
    from ..container import Document

__all__ = ["MONOSPACE_FONT", "write"]

logger = log.get_logger("document.export.docx")

#: Font a code element and a preformatted block are written in. Present on Windows, macOS
#: and most Linux installs, and a word processor substitutes another monospace font where it
#: is not.
MONOSPACE_FONT = "Consolas"

#: Colour and underline a link is drawn with, matching what Word's own Hyperlink style
#: gives. Set on the run rather than through that style, which a template need not carry.
LINK_COLOR = "0563C1"

#: Points one level of list nesting, quotation or definition indenting adds.
INDENT_STEP = 18.0

#: Deepest list nesting the built-in styles cover. A deeper item keeps the third style and
#: is indented further, so its text still sits under its parent.
STYLE_LEVELS = 3

#: Points of space left under a paragraph standing in for a horizontal rule.
RULE_SPACING = 6.0

#: Name of every built-in style used, so a template missing one is reported once by name.
_HEADING = "Heading {level}"
_BULLET = "List Bullet"
_NUMBER = "List Number"
_QUOTE = "Quote"
_CAPTION = "Caption"
_TABLE = "Table Grid"


def write(document: "Document", page: Page) -> bytes:
    """Write a document as a ``.docx`` file.

    Args:
        document: The document to convert.
        page: The page to lay it out on.

    Returns:
        The file's bytes.

    Raises:
        DependencyError: python-docx is missing or unusable.
    """
    deps.require("docx", feature=FEATURE)
    from docx import Document as WordDocument

    word = WordDocument()
    _page(word, page)
    _Writer(word, page).document(blocks.to_blocks(document.content, document.assets))
    _core_properties(word, document)
    buffer = BytesIO()
    word.save(buffer)
    return ooxml.with_custom_properties(buffer.getvalue(), _custom_properties(document))


def _page(word, page: Page) -> None:
    """Set the page size, orientation and margins on the document's one section."""
    from docx.enum.section import WD_ORIENT
    from docx.shared import Pt

    section = word.sections[0]
    section.page_width = Pt(page.width)
    section.page_height = Pt(page.height)
    section.orientation = (
        WD_ORIENT.LANDSCAPE if page.width > page.height else WD_ORIENT.PORTRAIT
    )
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, side, Pt(page.margin))


class _Writer:
    """Writes blocks into one python-docx document.

    Attributes:
        word: The document being built.
        page: The page it is laid out on.
    """

    def __init__(self, word, page: Page) -> None:
        """Write into one document.

        Args:
            word: The python-docx ``Document`` to write into.
            page: The page it is laid out on.
        """
        self.word = word
        self.page = page
        self._missing: set[str] = set()
        self._dropped = 0

    def document(self, found: list[blocks.Block]) -> None:
        """Write every block, then report anything that could not be drawn."""
        for block in found:
            self.block(block)
        if self._dropped:
            logger.warning(
                "%d picture(s) could not be drawn in the exported .docx and its description "
                "was written in its place. A picture is drawn when the document carries the "
                "file itself in a format Word reads: PNG, JPEG, GIF, BMP or TIFF. One named "
                "by a web address is never fetched.",
                self._dropped,
            )

    def block(self, block: blocks.Block) -> None:
        """Write one block."""
        if block.kind == "table":
            self._table(block)
            return
        if block.kind == "rule":
            self._rule(block)
            return
        if block.kind == "preformatted":
            self._preformatted(block)
            return
        paragraph = self._paragraph(block)
        self._runs(paragraph, block.runs)

    # ----------------------------------------------------------- the paragraph

    def _paragraph(self, block: blocks.Block, cell=None):
        """Start a paragraph for one block, in a table cell or in the body."""
        style = self._style_name(block)
        paragraph = self._new(cell, style)
        self._format(paragraph, block)
        return paragraph

    def _new(self, cell, style: str | None):
        """A new paragraph, or a cell's own first one where it is still empty."""
        if cell is None:
            return self._styled(self.word.add_paragraph(), style)
        first = cell.paragraphs[0]
        if not first.runs and not first.text:
            return self._styled(first, style)
        return self._styled(cell.add_paragraph(), style)

    def _styled(self, paragraph, style: str | None):
        """Apply a built-in style, leaving the paragraph as it is where none exists."""
        if not style:
            return paragraph
        try:
            paragraph.style = self.word.styles[style]
        except KeyError:
            if style not in self._missing:
                self._missing.add(style)
                logger.debug(
                    "this python-docx template carries no %r style, so those paragraphs are "
                    "written in the default one", style,
                )
        return paragraph

    def _style_name(self, block: blocks.Block) -> str | None:
        """The built-in style one block is written in, or ``None`` for the default."""
        if block.kind == "heading":
            return _HEADING.format(level=max(1, min(block.level, 9)))
        if block.kind == "item":
            base = _NUMBER if block.ordered else _BULLET
            level = max(1, min(block.level, STYLE_LEVELS))
            return base if level == 1 else f"{base} {level}"
        if block.quote:
            return _QUOTE
        return None

    def _format(self, paragraph, block: blocks.Block) -> None:
        """Set alignment and indenting on a paragraph."""
        from docx.shared import Pt

        alignment = _alignment(block.align)
        if alignment is not None:
            paragraph.alignment = alignment
        steps = block.indent + block.quote
        if block.kind == "item":
            steps += max(0, block.level - STYLE_LEVELS)
        if steps:
            paragraph.paragraph_format.left_indent = Pt(INDENT_STEP * steps)

    # ---------------------------------------------------------------- the runs

    def _runs(self, paragraph, found: list[blocks.Run]) -> None:
        """Write every run of one block into a paragraph."""
        for run in found:
            if run.line_break:
                paragraph.add_run().add_break()
            elif run.picture is not None:
                self._picture(paragraph, run)
            elif run.text:
                self._text(paragraph, run)

    def _text(self, paragraph, run: blocks.Run) -> None:
        """Write one run of text with its formatting."""
        from docx.shared import Pt, RGBColor

        style = run.style
        word_run = _hyperlink(paragraph, style.link, run.text) if style.link else None
        if word_run is None:
            word_run = paragraph.add_run(run.text)
        word_run.bold = style.bold or None
        word_run.italic = style.italic or None
        word_run.underline = True if style.underline or style.link else None
        font = word_run.font
        if style.strike:
            font.strike = True
        if style.superscript:
            font.superscript = True
        if style.subscript:
            font.subscript = True
        name = style.font or (MONOSPACE_FONT if style.monospace else None)
        if name:
            font.name = name
        if style.size:
            font.size = Pt(style.size)
        color = style.color or (LINK_COLOR if style.link else None)
        if color:
            font.color.rgb = RGBColor.from_string(color)
        if style.background:
            _shading(word_run, style.background)

    def _picture(self, paragraph, run: blocks.Run) -> None:
        """Draw a picture, or write its description where it cannot be drawn."""
        from docx.shared import Pt

        picture = run.picture
        if picture.data:
            width = Pt(min(picture.width, self.page.text_width)) if picture.width else None
            height = Pt(picture.height) if picture.height else None
            try:
                shape = paragraph.add_run().add_picture(
                    BytesIO(picture.data), width=width, height=height
                )
            except Exception as error:
                # Every reason a picture will not go in reaches the reader the same way, and
                # python-docx raises its own exception types for an unreadable image.
                logger.debug("%r could not be drawn (%s)", picture.source, error)
            else:
                _fit(shape, Pt(self.page.text_width))
                return
        else:
            logger.debug(
                "%r holds no picture this document carries%s",
                picture.source, ", and a web address is never fetched" if picture.remote else "",
            )
        self._dropped += 1
        placeholder = picture.alt or picture.source or "picture"
        run_with_alt = paragraph.add_run(f"[{placeholder}]")
        run_with_alt.italic = True

    # -------------------------------------------------------------- the blocks

    def _preformatted(self, block: blocks.Block) -> None:
        """Write a preformatted block, one paragraph per line, spaces kept."""
        from docx.shared import Pt

        for line in block.text.split("\n"):
            paragraph = self.word.add_paragraph()
            self._format(paragraph, block)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(line)
            run.font.name = MONOSPACE_FONT
            _preserve_space(run)

    def _rule(self, block: blocks.Block) -> None:
        """Write a horizontal rule as a paragraph carrying a bottom border."""
        from docx.shared import Pt

        paragraph = self.word.add_paragraph()
        self._format(paragraph, block)
        paragraph.paragraph_format.space_after = Pt(RULE_SPACING)
        _bottom_border(paragraph)

    def _table(self, block: blocks.Block) -> None:
        """Write a table, merging every cell that spans more than one position."""
        table = block.table
        matrix = blocks.grid(table)
        if not matrix or not matrix[0]:
            return
        if table.caption:
            caption = self._new(None, _CAPTION)
            self._runs(caption, table.caption)
        word_table = self.word.add_table(rows=len(matrix), cols=len(matrix[0]))
        try:
            word_table.style = self.word.styles[_TABLE]
        except KeyError:
            if _TABLE not in self._missing:
                self._missing.add(_TABLE)
                logger.debug(
                    "this python-docx template carries no %r style, so the exported table "
                    "has no cell borders", _TABLE,
                )
        for row, slots in enumerate(matrix):
            for column, slot in enumerate(slots):
                if slot.origin and (slot.colspan > 1 or slot.rowspan > 1):
                    _merge(word_table, row, column, slot)
        for row, slots in enumerate(matrix):
            for column, slot in enumerate(slots):
                if slot.origin:
                    self._cell(word_table.cell(row, column), slot.cell)

    def _cell(self, word_cell, cell: blocks.Cell) -> None:
        """Write one cell's blocks into it."""
        for block in cell.blocks or [blocks.Block()]:
            if block.kind == "table":
                # A table nested inside a cell is written as its rows of text, one paragraph
                # per row: python-docx nests a table only in a cell it is added to directly,
                # and rebuilding the nesting here would lose the spans either way.
                for row in block.table.rows if block.table else []:
                    paragraph = self._paragraph(blocks.Block(align=cell.align), word_cell)
                    for index, inner in enumerate(row.cells):
                        if index:
                            paragraph.add_run("\t")
                        for nested in inner.blocks:
                            self._runs(paragraph, nested.runs)
                continue
            if block.kind == "rule":
                _bottom_border(self._paragraph(blocks.Block(), word_cell))
                continue
            if block.kind == "preformatted":
                paragraph = self._paragraph(blocks.Block(align=cell.align), word_cell)
                run = paragraph.add_run(block.text)
                run.font.name = MONOSPACE_FONT
                _preserve_space(run)
                continue
            if block.align is None and cell.align:
                block = replace(block, align=cell.align)
            self._runs(self._paragraph(block, word_cell), block.runs)


def _alignment(align: str | None):
    """One CSS alignment as a python-docx paragraph alignment, or ``None``."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align or "")


def _fit(shape, limit) -> None:
    """Scale an inline shape down to the text column, keeping its proportions."""
    try:
        if shape.width and shape.width > limit:
            height = int(shape.height * limit / shape.width)
            shape.width = int(limit)
            shape.height = height
    except (TypeError, ValueError, ZeroDivisionError) as error:
        logger.debug("a picture could not be scaled to the page (%s)", error)


def _merge(table, row: int, column: int, slot: blocks.Slot) -> None:
    """Merge the rectangle one spanning cell covers."""
    try:
        first = table.cell(row, column)
        last = table.cell(row + slot.rowspan - 1, column + slot.colspan - 1)
        if first is not last:
            first.merge(last)
    except (IndexError, ValueError, KeyError) as error:
        logger.debug(
            "a cell spanning %d column(s) and %d row(s) was written unmerged (%s)",
            slot.colspan, slot.rowspan, error,
        )


def _hyperlink(paragraph, url: str, text: str):
    """Add a real hyperlink to a paragraph.

    python-docx has no hyperlink of its own, so the element and its package relationship are
    written directly.

    Args:
        paragraph: The paragraph the link is added to.
        url: Where it points.
        text: The words it is drawn as.

    Returns:
        The run holding the text, or ``None`` where the link could not be written, leaving
        the caller to add ordinary text instead.
    """
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.text.run import Run

        relationship = paragraph.part.relate_to(
            url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        anchor = OxmlElement("w:hyperlink")
        anchor.set(qn("r:id"), relationship)
        element = OxmlElement("w:r")
        anchor.append(element)
        paragraph._p.append(anchor)
        run = Run(element, paragraph)
        run.text = text
        return run
    except Exception as error:
        # A link that cannot be written is worth less than the words it was written on, so
        # the words are kept and the address is the only loss.
        logger.debug("the link to %r was written as plain text (%s)", url, error)
        return None


def _shading(run, color: str) -> None:
    """Fill the background behind one run with a colour."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color)
        run._element.get_or_add_rPr().append(shading)
    except Exception as error:
        logger.debug("a background colour was not written (%s)", error)


def _bottom_border(paragraph) -> None:
    """Draw a line under a paragraph, which is how a rule is written."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        borders.append(bottom)
        paragraph._p.get_or_add_pPr().append(borders)
    except Exception as error:
        logger.debug("a horizontal rule was written as an empty paragraph (%s)", error)


def _preserve_space(run) -> None:
    """Keep the leading and trailing spaces of a run, which a reader would drop."""
    try:
        from docx.oxml.ns import qn

        for element in run._element.findall(qn("w:t")):
            element.set(qn("xml:space"), "preserve")
    except Exception as error:
        logger.debug("the spacing of a preformatted line may be lost (%s)", error)


def _core_properties(word, document: "Document") -> None:
    """Write the metadata OOXML has a core property for."""
    metadata = document.metadata
    properties = word.core_properties
    properties.title = metadata.title
    properties.author = metadata.author
    properties.comments = metadata.description
    properties.keywords = keywords_text(metadata.keywords)
    properties.language = metadata.language
    created = _stamp(metadata.created)
    modified = _stamp(metadata.modified)
    if created:
        properties.created = created
    if modified:
        properties.modified = modified


def _custom_properties(document: "Document") -> dict[str, str]:
    """The metadata OOXML has no core property for, as custom properties."""
    metadata = document.metadata
    found: dict[str, str] = {}
    if metadata.copyright:
        found["Copyright"] = metadata.copyright
    if metadata.generator:
        found["Generator"] = metadata.generator
    for name, value in metadata.custom.items():
        if name not in found:
            found[name] = value
    return found


def _stamp(value: str) -> Any:
    """One document timestamp as a datetime, or ``None`` where it is not one."""
    if not value:
        return None
    try:
        return datetime.strptime(value, metadata_module.STAMP_FORMAT).replace(
            tzinfo=timezone.utc
        )
    except ValueError:
        logger.debug("%r is not a timestamp, so it was left out of the exported file", value)
        return None
