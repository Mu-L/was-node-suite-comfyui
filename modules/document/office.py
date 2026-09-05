"""Reading a ``.docx`` or an ``.odt`` file into a document.

:func:`read` answers a :class:`.container.Document`: the file's writing as HTML, its
metadata, and each picture it carries as an embedded file named relative to ``assets/``.
"""

from __future__ import annotations

import base64
import binascii
import html as html_escaping
import re
import zipfile
from dataclasses import replace
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Mapping
from xml.etree import ElementTree

from .. import deps, log
from . import container, markup
from .export import DOCX, FEATURE, ODT
from .export import PACKAGES as WRITER_PACKAGES
from .export import css, ooxml
from .export.blocks import Style
from .metadata import STAMP_FORMAT, Metadata

__all__ = [
    "DOCX",
    "FEATURE",
    "FORMATS",
    "FORMAT_NAMES",
    "MONOSPACE_FONTS",
    "ODT",
    "PACKAGES",
    "normalized",
    "read",
    "reads",
]

logger = log.get_logger("document.office")

#: The two formats read here, each spelled as the extension its file carries.
FORMATS = (DOCX, ODT)

#: Import name of the package each format needs, the same one that writes it.
PACKAGES = {extension: WRITER_PACKAGES[extension] for extension in FORMATS}

#: A part every file of that format holds, so a zip renamed to one is named as what it is.
REQUIRED_ENTRY = {DOCX: "word/document.xml", ODT: "content.xml"}

#: What each format is called in a message.
FORMAT_NAMES = {DOCX: "Word document", ODT: "OpenDocument text document"}

#: Fonts read back as a code element rather than as a font name: the two the export writes,
#: and the monospace families a word processor offers beside them.
MONOSPACE_FONTS = frozenset(
    {
        "consolas", "courier", "courier new", "dejavu sans mono", "liberation mono",
        "lucida console", "menlo", "monaco", "monospace", "nimbus mono ps",
        "sf mono", "source code pro", "ubuntu mono",
    }
)

#: Deepest heading a document carries. HTML has six levels and both formats offer more.
MAX_HEADING = 6

#: Longest run of spaces one ODF space element stands for, most times a table row or cell
#: may repeat, and how far one cell may span. Each is an attribute, so a file may claim
#: millions of any of them.
MAX_SPACES = 512
MAX_REPEAT = 64
MAX_SPAN = 64

#: How far a style may inherit before the chain is read as a loop.
MAX_STYLE_DEPTH = 16

#: English Metric Units in one point, and twips in one point: the two units a ``.docx``
#: measures a picture and an indent in.
EMU_PER_POINT = 12700.0
TWIPS_PER_POINT = 20.0

#: Where a ``.docx`` names the program that wrote it.
APP_PART = "docProps/app.xml"

#: A picture drawn the older way, as a shape rather than a drawing. python-docx declares no
#: prefix for that namespace, so the element is named in full.
VML_IMAGE = "{urn:schemas-microsoft-com:vml}imagedata"

#: Alignments OOXML spells its own way, as CSS spells them.
DOCX_ALIGNMENTS = {"both": "justify", "distribute": "justify"}

#: Media type at the front of every OpenDocument text package.
ODT_MIMETYPE = "application/vnd.oasis.opendocument.text"

#: ODF elements read through to reach the blocks inside them.
ODT_GROUPS = (
    "office:text", "text:index-body", "text:index-title", "text:section",
    "text:table-of-content",
)

#: ODF elements that hold nothing a reader sees, and are passed over without a report.
ODT_IGNORED = (
    "office:forms", "table:table-column", "table:table-columns", "text:sequence-decls",
    "text:soft-page-break", "text:tracked-changes", "text:user-field-decls",
    "text:variable-decls",
)

#: ODF elements inside a paragraph that mark a place rather than write anything.
ODT_SKIPPED_INLINE = (
    "office:annotation-end", "text:bookmark", "text:bookmark-end", "text:bookmark-start",
    "text:reference-mark", "text:reference-mark-end", "text:reference-mark-start",
    "text:soft-page-break",
)

#: ODF elements inside a paragraph whose content a document has no place for, and what each
#: is called when the count is reported.
ODT_DROPPED_INLINE = {
    "office:annotation": "comment(s)",
    "text:note": "footnote(s) and endnote(s)",
}

#: ODF styles whose properties are the document defaults. Each one keeps its place in a
#: style chain, so a block is still recognised by it, and none of them formats a run.
ODF_BASE_STYLES = frozenset(
    {
        "caption", "default paragraph font", "default style", "endnote", "footnote",
        "heading", "illustration", "index", "list", "list contents", "preformatted text",
        "quotations", "standard", "table contents", "table heading", "text body",
    }
)

#: Extension each picture format is stored under, by the bytes the file starts with.
SIGNATURES = (
    (b"\x89PNG\r\n\x1a\n", ".png"),
    (b"\xff\xd8\xff", ".jpg"),
    (b"GIF87a", ".gif"),
    (b"GIF89a", ".gif"),
    (b"BM", ".bmp"),
    (b"II*\x00", ".tif"),
    (b"MM\x00*", ".tif"),
    (b"RIFF", ".webp"),
    (b"<?xml", ".svg"),
    (b"<svg", ".svg"),
)

#: Extension a picture is stored under when its bytes name no format.
DEFAULT_PICTURE_SUFFIX = ".img"

#: Schemes a picture may name that sit on the network. Nothing here fetches one.
REMOTE_SCHEMES = ("http://", "https://", "ftp://", "//")

#: Character formatting to the element it is written back as, innermost first.
_INLINE_ELEMENTS = (
    ("monospace", "code"),
    ("subscript", "sub"),
    ("superscript", "sup"),
    ("strike", "s"),
    ("underline", "u"),
    ("italic", "em"),
    ("bold", "strong"),
)

#: A cell holding one unstyled paragraph, whose tags come off so its text sits in the cell.
_LONE_PARAGRAPH = re.compile(r"<p>(.*)</p>", re.DOTALL)

#: A document type declaration, which is not read out of a metadata part.
_DOCTYPE = re.compile(rb"<!DOCTYPE", re.IGNORECASE)

#: How a Word style names a heading level, and how it names a list and its level.
_HEADING_STYLE = re.compile(r"^heading\s*(\d+)$")
_LIST_STYLE = re.compile(r"^list\s+(bullet|number)\s*(\d*)$")

#: How ODF spells a space inside a style name.
ODF_SPACE = "_20_"

#: ODF style names that say what kind of block a paragraph is.
ODF_QUOTE = ("quotations", "quote")
ODF_PREFORMATTED = ("preformatted text",)
ODF_CAPTION = ("illustration", "caption", "figure")


def reads(extension: Any) -> bool:
    """Whether one format is read here rather than by :mod:`.container`.

    Args:
        extension: A file suffix, with or without its leading dot and in any case.

    Returns:
        True for one of :data:`FORMATS`.
    """
    return _spelled(extension) in FORMATS


def normalized(extension: Any) -> str:
    """One of :data:`FORMATS`, read from a file suffix.

    Args:
        extension: The suffix. A leading dot is added where it is missing and the case is
            ignored, so ``DOCX`` and ``.docx`` both name the same format.

    Returns:
        The format, spelled as :data:`FORMATS` spells it.

    Raises:
        ValueError: The value names neither format.
    """
    text = _spelled(extension)
    if text not in FORMATS:
        raise ValueError(
            f"{extension!r} is not one of the formats read through a document library.\n"
            f"  Those formats are: {', '.join(FORMATS)}."
        )
    return text


def read(data: Any, extension: Any) -> container.Document:
    """Read one ``.docx`` or ``.odt`` file.

    Args:
        data: The file's bytes.
        extension: One of :data:`FORMATS`.

    Returns:
        The document: the file's writing as HTML, the metadata it carries, and every
        picture stored inside it as an embedded file.

    Raises:
        NotADocument: The bytes are not a file of that format.
        DocumentError: The file unpacks past the bounds a document is read under, or the
            markup built from it cannot go in a container.
        ValueError: ``extension`` names neither format.
        DependencyError: The package that format needs is missing or unusable.
    """
    chosen = normalized(extension)
    payload = _as_bytes(data, chosen)
    _bounded(payload, chosen)
    document = _read_docx(payload) if chosen == DOCX else _read_odt(payload)
    logger.info(
        "read %s: %d word(s), %d character(s), %d picture(s)",
        _a(chosen), document.word_count, document.character_count, len(document.assets),
    )
    return document


# --------------------------------------------------------------------- the package


def _a(extension: str) -> str:
    """The name of one format with the article that belongs in front of it."""
    name = FORMAT_NAMES[extension]
    return f"{'an' if name[:1].upper() in 'AEIOU' else 'a'} {name}"


def _spelled(extension: Any) -> str:
    """One suffix as a format is spelled here: lower case, leading dot."""
    text = str(extension or "").strip().lower()
    if text and not text.startswith("."):
        text = "." + text
    return text


def _as_bytes(data: Any, extension: str) -> bytes:
    """The file's bytes, copied where the original could still be written to.

    Args:
        data: The value offered.
        extension: The format, for the message.

    Returns:
        Immutable bytes.

    Raises:
        NotADocument: ``data`` is not bytes at all.
    """
    if isinstance(data, (bytearray, memoryview)):
        return bytes(data)
    if isinstance(data, bytes):
        return data
    name = type(data).__name__
    raise container.NotADocument(
        f"a {extension} file is read from bytes, and this is "
        f"{'an' if name[:1].lower() in 'aeiou' else 'a'} {name}."
    )


def _bounded(payload: bytes, extension: str) -> None:
    """Confirm the bytes are that format's package and fit what a document may unpack to.

    Args:
        payload: The file's bytes.
        extension: One of :data:`FORMATS`.

    Raises:
        NotADocument: The bytes are not a readable zip, or hold none of that format's parts.
        DocumentError: The package declares more than
            :data:`.container.MAX_UNPACKED_BYTES`.
    """
    try:
        with zipfile.ZipFile(BytesIO(payload)) as archive:
            infos = archive.infolist()
    except (zipfile.BadZipFile, OSError, ValueError) as error:
        raise container.NotADocument(
            f"this is not a {extension} file: one is a zip package, and these "
            f"{len(payload)} byte(s) are not a zip file ({error}).\n"
            f"  Every {FORMAT_NAMES[extension]} starts with the two bytes 'PK', as every "
            f"zip file does. A file written in the older binary .doc format is not one of "
            f"these: open it in a word processor and save it as {extension} first."
        ) from error
    declared = sum(max(int(info.file_size), 0) for info in infos)
    if declared > container.MAX_UNPACKED_BYTES:
        raise container.DocumentError(
            f"refusing to read a {extension} file that unpacks to "
            f"{declared / (1024 * 1024):.1f} MB; the limit is "
            f"{container.MAX_UNPACKED_BYTES / (1024 * 1024):.1f} MB.\n"
            f"  A file that large is either damaged or built to exhaust memory."
        )
    names = [info.filename for info in infos if not info.is_dir()]
    required = REQUIRED_ENTRY[extension]
    if required not in names:
        listed = ", ".join(sorted(names)[:12]) or "nothing"
        raise container.NotADocument(
            f"this zip file is not {_a(extension)}: every {extension} package holds a "
            f"{required} part, and this one holds {listed}.\n"
            f"  To read every document inside an archive, use Load Documents from ZIP."
        )


# --------------------------------------------------------------------- the markup


def _text(value: str) -> str:
    """One piece of writing with the three characters markup gives a meaning to escaped."""
    return html_escaping.escape(str(value), quote=False)


def _attribute(value: str) -> str:
    """One attribute value, quotes escaped as well."""
    return html_escaping.escape(str(value), quote=True)


def _inline(text: str, style: Style) -> str:
    """One run of text wrapped in the elements its formatting needs.

    Args:
        text: The characters, escaped here.
        style: The formatting in force.

    Returns:
        The markup, or an empty string where there is no text.
    """
    if not text:
        return ""
    body = _text(text)
    declarations = []
    if style.color:
        declarations.append(f"color:#{style.color}")
    if style.background:
        declarations.append(f"background-color:#{style.background}")
    if style.font and not style.monospace:
        declarations.append(f"font-family:{style.font}")
    if style.size:
        declarations.append(f"font-size:{style.size:g}pt")
    if declarations:
        body = f'<span style="{_attribute(";".join(declarations))}">{body}</span>'
    for flag, element in _INLINE_ELEMENTS:
        if getattr(style, flag):
            body = f"<{element}>{body}</{element}>"
    if style.link:
        body = f'<a href="{_attribute(style.link)}">{body}</a>'
    return body


def _font_name(value: Any) -> str | None:
    """One font family name, or ``None`` where the value names none."""
    name = css.font_family(str(value or ""))
    if not name:
        return None
    return name.replace(";", " ").replace('"', "").strip() or None


def _block_style(align: str | None, indent: float | None) -> str:
    """The ``style`` attribute one block carries, or an empty string."""
    declarations = []
    if align:
        declarations.append(f"text-align:{align}")
    if indent:
        declarations.append(f"margin-left:{indent:g}pt")
    return ";".join(declarations)


def _style_attribute(style: str) -> str:
    """One block style as an attribute, ready to sit inside a start tag."""
    return f' style="{_attribute(style)}"' if style else ""


def _picture_markup(source: str, alt: str, width: float | None, height: float | None) -> str:
    """One picture as an ``img`` element."""
    sizes = []
    if width:
        sizes.append(f"width:{width:g}pt")
    if height:
        sizes.append(f"height:{height:g}pt")
    return (
        f'<img src="{_attribute(source)}" alt="{_attribute(alt)}"'
        f"{_style_attribute(';'.join(sizes))}>"
    )


class _Runs:
    """Text collected with the formatting it was written in, wrapped as the style changes."""

    def __init__(self) -> None:
        """Start with nothing collected."""
        self._parts: list[str] = []
        self._text = ""
        self._style = Style()

    def text(self, value: str, style: Style) -> None:
        """Add characters written in one style, joining them to the run before."""
        if not value:
            return
        if self._text and style != self._style:
            self.flush()
        self._style = style
        self._text += value

    def element(self, value: str) -> None:
        """Add markup that is not text, such as a line break or a picture."""
        self.flush()
        if value:
            self._parts.append(value)

    def flush(self) -> None:
        """Wrap and keep whatever text has been collected."""
        if self._text:
            self._parts.append(_inline(self._text, self._style))
        self._text = ""
        self._style = Style()

    def result(self) -> str:
        """Everything collected, as the markup of one block."""
        self.flush()
        return "".join(self._parts)


class _Markup:
    """HTML assembled from one office file, block by block, in the order it was written."""

    def __init__(self) -> None:
        """Start with nothing written."""
        self._parts: list[str] = []
        self._lists: list[bool] = []
        self._item_open = False
        self._quote_open = False
        self._caption = ""

    def heading(self, level: int, inner: str, style: str = "") -> None:
        """Write one heading, its level clamped to what HTML offers."""
        self._before(False)
        depth = max(1, min(int(level or 1), MAX_HEADING))
        self._line(f"<h{depth}{_style_attribute(style)}>{inner}</h{depth}>")

    def paragraph(self, inner: str, style: str = "", quote: bool = False) -> None:
        """Write one paragraph, inside a quotation where it belongs to one."""
        self._before(quote)
        self._line(f"<p{_style_attribute(style)}>{inner}</p>")

    def preformatted(self, inner: str, style: str = "") -> None:
        """Write one block whose spacing and line breaks are part of the writing."""
        self._before(False)
        self._line(f"<pre{_style_attribute(style)}>{inner}</pre>")

    def rule(self, style: str = "") -> None:
        """Write a horizontal rule."""
        self._before(False)
        self._line(f"<hr{_style_attribute(style)}>")

    def caption(self, inner: str) -> None:
        """Hold a caption for the table that may follow it."""
        self._before(False)
        self._caption = inner

    def table(self, rows: str, columns: str = "") -> None:
        """Write one table, taking any caption held for it."""
        self._close_lists()
        self._quote(False)
        caption = f"<caption>{self._caption}</caption>" if self._caption else ""
        self._caption = ""
        self._line(f"<table>{caption}{columns}{rows}</table>")

    def item(self, level: int, ordered: bool, inner: str, style: str = "") -> None:
        """Write one list item, opening and closing lists as the level changes."""
        self._flush_caption()
        self._quote(False)
        depth = max(1, int(level or 1))
        while len(self._lists) > depth:
            self._pop()
        if self._lists and len(self._lists) == depth and self._lists[-1] != ordered:
            self._pop()
            self._item_open = False
        while len(self._lists) < depth:
            if self._lists and not self._item_open:
                self._parts.append("<li>")
                self._item_open = True
            self._line("<ol>" if ordered else "<ul>")
            self._lists.append(ordered)
            self._item_open = False
        if self._item_open:
            self._line("</li>")
        self._parts.append(f"<li{_style_attribute(style)}>{inner}")
        self._item_open = True

    def result(self) -> str:
        """Everything written, with every element closed.

        Returns:
            The document markup, one block to a line.
        """
        self._flush_caption()
        self._close_lists()
        self._quote(False)
        return "".join(self._parts).strip()

    def _line(self, part: str) -> None:
        """Write one block on a line of its own."""
        self._parts.append(part + "\n")

    def _before(self, quote: bool) -> None:
        """Close whatever the next block cannot sit inside."""
        self._flush_caption()
        self._close_lists()
        self._quote(quote)

    def _flush_caption(self) -> None:
        """Write a held caption as a paragraph, no table having followed it."""
        if self._caption:
            held, self._caption = self._caption, ""
            self._line(f"<p>{held}</p>")

    def _quote(self, wanted: bool) -> None:
        """Open or close the quotation the following blocks sit in."""
        if wanted and not self._quote_open:
            self._line("<blockquote>")
            self._quote_open = True
        elif not wanted and self._quote_open:
            self._line("</blockquote>")
            self._quote_open = False

    def _pop(self) -> None:
        """Close the innermost list, leaving the item that held it open."""
        if self._item_open:
            self._line("</li>")
        self._line("</ol>" if self._lists.pop() else "</ul>")
        self._item_open = True

    def _close_lists(self) -> None:
        """Close every open list."""
        while self._lists:
            self._pop()
        self._item_open = False


def _cell(html: str, header: bool, colspan: int, rowspan: int) -> str:
    """One table cell, its lone paragraph unwrapped so the text sits in the cell."""
    match = _LONE_PARAGRAPH.fullmatch(html.strip())
    if match and "<p" not in match.group(1) and "<table" not in match.group(1):
        html = match.group(1)
    tag = "th" if header else "td"
    spans = ""
    if colspan > 1:
        spans += f' colspan="{colspan}"'
    if rowspan > 1:
        spans += f' rowspan="{rowspan}"'
    return f"<{tag}{spans}>{html}</{tag}>"


# --------------------------------------------------------------------- the pictures


class _Assets:
    """The pictures pulled out of one file, named the way a document names them.

    Attributes:
        found: ``{name: bytes}``, ready to be a document's embedded files.
    """

    def __init__(self) -> None:
        """Start with no picture stored."""
        self.found: dict[str, bytes] = {}
        self._sources: dict[str, str | None] = {}
        self._bytes = 0
        self._refused = 0

    def add(self, source: str, data: bytes, suggested: str) -> str | None:
        """Store one picture and answer the name it is embedded under.

        Args:
            source: How the file refers to it, so the same picture drawn twice is stored
                once.
            data: The picture's bytes.
            suggested: The name it carries inside the file.

        Returns:
            The name under ``assets/``, or ``None`` when there was no room for it.
        """
        if source in self._sources:
            return self._sources[source]
        if not data:
            return None
        if (
            len(self.found) >= container.MAX_ASSETS
            or self._bytes + len(data) > container.MAX_UNPACKED_BYTES
        ):
            self._refused += 1
            self._sources[source] = None
            return None
        name = _unique(_picture_name(suggested, data, len(self.found)), self.found)
        self.found[name] = data
        self._bytes += len(data)
        self._sources[source] = name
        return name

    def report(self, extension: str) -> None:
        """Log every picture there was no room for."""
        if self._refused:
            logger.warning(
                "%d picture(s) in this %s file were left out of the document: it already "
                "carries %d embedded file(s) and %.1f MB of them, which is as much as a "
                "document holds. The writing around them was read in full, and each one is "
                "now a broken link.",
                self._refused, extension, len(self.found), self._bytes / (1024 * 1024),
            )


def _picture_name(suggested: str, data: bytes, index: int) -> str:
    """A name for one picture inside the document, with the extension its bytes call for."""
    stem = str(suggested or "").replace("\\", "/").rsplit("/", 1)[-1].strip()
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._") or f"image{index + 1}"
    known = _suffix(data)
    if known and not stem.lower().endswith(known):
        stem = re.sub(r"\.[A-Za-z0-9]{1,5}$", "", stem) + known
    elif "." not in stem:
        stem += DEFAULT_PICTURE_SUFFIX
    return stem


def _suffix(data: bytes) -> str:
    """The extension one picture's bytes name, or an empty string."""
    opening = bytes(data[:16])
    for signature, extension in SIGNATURES:
        if opening.startswith(signature):
            return extension
    return ""


def _unique(name: str, taken: Mapping[str, bytes]) -> str:
    """The name itself, or the next one like it that nothing already holds."""
    if name not in taken:
        return name
    stem, dot, extension = name.rpartition(".")
    stem = stem or name
    extension = f"{dot}{extension}" if dot else ""
    number = 2
    while f"{stem}-{number}{extension}" in taken:
        number += 1
    return f"{stem}-{number}{extension}"


def _remote(source: str) -> bool:
    """Whether a picture's address sits on the network."""
    return str(source or "").strip().lower().startswith(REMOTE_SCHEMES)


# --------------------------------------------------------------------- the metadata


def _stamp(value: Any) -> str:
    """One date out of a file as a document timestamp, or an empty string."""
    if isinstance(value, datetime):
        moment = value if value.tzinfo else value.replace(tzinfo=timezone.utc)
        return moment.astimezone(timezone.utc).strftime(STAMP_FORMAT)
    text = str(value or "").strip()
    if not text:
        return ""
    try:
        return _stamp(datetime.fromisoformat(text.replace("Z", "+00:00")))
    except ValueError:
        logger.debug("%r is not a date, so it was left out of the document", text)
        return ""


def _parsed(payload: bytes, part: str):
    """One metadata part as an element tree, or ``None`` where it cannot be read.

    Args:
        payload: The part's bytes.
        part: Its name inside the package, for the message.

    Returns:
        The root element, or ``None``.
    """
    if _DOCTYPE.search(payload[:4096]):
        logger.warning(
            "%s declares a document type, which can define entities that expand to far "
            "more than the file holds, so it was not read. The document carries whatever "
            "the rest of the file states and none of the properties this part held.",
            part,
        )
        return None
    try:
        return ElementTree.fromstring(payload)
    except (ElementTree.ParseError, ValueError) as error:
        logger.debug("%s could not be read (%s)", part, error)
        return None


def _cleaned(content: str, extension: str) -> str:
    """The markup with script, frame and event-handler content taken out of it."""
    result, removed = markup.clean(content)
    if removed:
        logger.info(
            "%s was removed from the %s file while it was read, and the writing around it "
            "was kept.",
            markup.describe(removed), extension,
        )
    return result


# ------------------------------------------------------------------------- .docx


def _read_docx(payload: bytes) -> container.Document:
    """Read a ``.docx`` package through python-docx.

    Args:
        payload: The package bytes.

    Returns:
        The document.

    Raises:
        NotADocument: python-docx would not open the package.
        DependencyError: python-docx is missing or unusable.
    """
    deps.require("docx", feature=FEATURE)
    from docx import Document as WordDocument

    try:
        word = WordDocument(BytesIO(payload))
    except Exception as error:
        # python-docx raises its own exception types for a package it cannot open.
        raise container.NotADocument(
            f"this {DOCX} file could not be opened by python-docx ({error}).\n"
            f"  The package is damaged, or it is a template or a macro-enabled document "
            f"rather than a plain {DOCX}. Open it in a word processor and save it again "
            f"as {DOCX}."
        ) from error
    reader = _Docx(word)
    reader.body()
    reader.assets.report(DOCX)
    reader.report()
    return container.Document(
        _cleaned(reader.page.result(), DOCX),
        _docx_metadata(word, payload),
        reader.assets.found,
    )


class _Docx:
    """Reads one python-docx document into markup and pictures.

    Attributes:
        word: The document being read.
        part: The package part its pictures and links hang from.
        page: The markup being assembled.
        assets: The pictures found so far.
    """

    def __init__(self, word) -> None:
        """Read one python-docx document.

        Args:
            word: The opened ``Document``.
        """
        from docx.oxml.ns import qn

        self.word = word
        self.part = word.part
        self.page = _Markup()
        self.assets = _Assets()
        self._qn = qn
        self._numbering = _numbering_formats(word, qn)
        self._through = {
            qn(name)
            for name in ("w:ins", "w:smartTag", "w:sdt", "w:sdtContent", "w:fldSimple")
        }
        self._dropped: dict[str, int] = {}

    def body(self) -> None:
        """Read every block of the document body."""
        for element in self.word.element.body.iterchildren():
            self.block(element)

    def report(self) -> None:
        """Log whatever the file holds that a document has no place for."""
        for what, count in sorted(self._dropped.items()):
            logger.info("%d %s in this %s file were not read.", count, what, DOCX)

    def block(self, element) -> None:
        """Read one element of a body, a table cell or a content control."""
        tag = element.tag
        if tag == self._qn("w:p"):
            self.paragraph(element)
        elif tag == self._qn("w:tbl"):
            self.table(element)
        elif tag in (self._qn("w:sdt"), self._qn("w:sdtContent")):
            for child in element.iterchildren():
                self.block(child)

    def paragraph(self, element) -> None:
        """Read one paragraph as whatever kind of block its style makes it."""
        from docx.text.paragraph import Paragraph

        name = _docx_style_name(Paragraph(element, self.word))
        runs = _Runs()
        self.content(element, Style(), runs)
        inner = runs.result()
        style = _block_style(
            _docx_align(element, self._qn), _docx_indent(element, self._qn)
        )
        level, ordered = self._listing(element, name)
        if level:
            self.page.item(level, ordered, inner, style)
            return
        heading = _docx_heading(name)
        if heading:
            self.page.heading(heading, inner, style)
            return
        if not inner.strip() and _docx_rule(element, self._qn):
            self.page.rule(style)
            return
        lowered = name.strip().lower()
        if lowered == "caption":
            self.page.caption(inner)
            return
        quoted = "quote" in lowered or "quotation" in lowered
        self.page.paragraph(inner, style, quote=quoted)

    def content(self, element, style: Style, runs: _Runs) -> None:
        """Read the runs of one paragraph, or of one element wrapped around them."""
        for child in element.iterchildren():
            tag = child.tag
            if tag == self._qn("w:r"):
                self.run(child, style, runs)
            elif tag == self._qn("w:hyperlink"):
                self.content(child, replace(style, link=self._link(child)), runs)
            elif tag == self._qn("w:del"):
                self._drop("tracked deletion(s)")
            elif tag in self._through:
                self.content(child, style, runs)

    def run(self, element, style: Style, runs: _Runs) -> None:
        """Read one run: its text, its breaks and any picture drawn in it."""
        style = self._run_format(element, style)
        for child in element.iterchildren():
            tag = child.tag
            if tag == self._qn("w:t"):
                runs.text(child.text or "", style)
            elif tag in (self._qn("w:br"), self._qn("w:cr")):
                runs.element("<br>")
            elif tag == self._qn("w:tab"):
                runs.text("\t", style)
            elif tag == self._qn("w:noBreakHyphen"):
                runs.text("-", style)
            elif tag in (self._qn("w:drawing"), self._qn("w:pict"), self._qn("w:object")):
                self.pictures(child, style, runs)
            elif tag == self._qn("w:delText"):
                self._drop("tracked deletion(s)")

    def pictures(self, element, style: Style, runs: _Runs) -> None:
        """Read every picture drawn inside one drawing, shape or embedded object."""
        found = []
        for blip in element.iter(self._qn("a:blip")):
            found.append(blip.get(self._qn("r:embed")) or blip.get(self._qn("r:link")))
        if not found:
            for image in element.iter(VML_IMAGE):
                found.append(image.get(self._qn("r:id")))
        if not found:
            self._drop("drawing(s) that hold no picture")
            return
        width, height, alt = _docx_drawing(element, self._qn)
        for identifier in found:
            runs.element(self._picture(identifier, alt, width, height, style))

    def table(self, element) -> None:
        """Read one table, its merged cells rebuilt as spans."""
        rows = []
        origins: dict[int, list | None] = {}
        for row in element.findall(self._qn("w:tr")):
            # A row every span reaches into keeps its place, empty.
            rows.append(self._row(row, origins))
        if not any(rows):
            return
        self.page.table(
            "".join(
                "<tr>"
                + "".join(
                    _cell(html, header, colspan, rowspan)
                    for html, rowspan, colspan, header in cells
                )
                + "</tr>"
                for cells in rows
            )
        )

    def cell(self, element) -> str:
        """Read one cell into markup of its own."""
        outer = self.page
        self.page = _Markup()
        try:
            for child in element.iterchildren():
                self.block(child)
            return self.page.result()
        finally:
            self.page = outer

    def _row(self, row, origins: dict[int, list | None]) -> list[list]:
        """One table row, each vertically merged cell folded into the one above it."""
        qn = self._qn
        properties = row.find(qn("w:trPr"))
        header = properties is not None and bool(
            _docx_toggle(properties, qn("w:tblHeader"), qn("w:val"))
        )
        cells: list[list] = []
        column = 0
        for cell in row.findall(qn("w:tc")):
            colspan, merge = _docx_cell_spans(cell, qn)
            continued = (
                merge is not None and (merge.get(qn("w:val")) or "continue") != "restart"
            )
            above = origins.get(column)
            if continued and above is not None:
                above[1] += 1
            else:
                record = [self.cell(cell), 1, colspan, header]
                cells.append(record)
                origins[column] = record if merge is not None else None
            column += colspan
        return cells

    def _picture(
        self,
        identifier: str | None,
        alt: str,
        width: float | None,
        height: float | None,
        style: Style,
    ) -> str:
        """One picture, or its description where the package does not carry the file."""
        part = None
        if identifier:
            try:
                part = self.part.related_parts[identifier]
            except (AttributeError, KeyError, OSError, ValueError) as error:
                logger.debug("a picture is not in the package (%s)", error)
        if part is None:
            self._drop("picture(s) the package does not carry")
            return _inline(f"[{alt or 'picture'}]", replace(style, italic=True))
        name = self.assets.add(str(part.partname), part.blob, str(part.partname))
        if name is None:
            return _inline(f"[{alt or 'picture'}]", replace(style, italic=True))
        return _picture_markup(container.ASSET_PREFIX + name, alt, width, height)

    def _link(self, element) -> str | None:
        """Where one hyperlink points, or ``None``."""
        identifier = element.get(self._qn("r:id"))
        if identifier:
            try:
                return self.part.rels[identifier].target_ref
            except (AttributeError, KeyError, ValueError) as error:
                logger.debug("a hyperlink has no address in the package (%s)", error)
                return None
        anchor = element.get(self._qn("w:anchor"))
        return f"#{anchor}" if anchor else None

    def _listing(self, element, name: str) -> tuple[int, bool]:
        """The list level and marker one paragraph belongs to, or ``(0, False)``."""
        level, ordered = _docx_list_style(name)
        reference = _docx_numbering(element, self._qn)
        if reference is None:
            return (max(1, level), ordered) if level else (0, False)
        level = max(level, reference[1] + 1)
        marker = self._numbering.get(reference)
        if marker:
            ordered = marker != "bullet"
        return max(1, level), ordered

    def _run_format(self, element, style: Style) -> Style:
        """The formatting one run declares, over the formatting around it."""
        qn = self._qn
        properties = element.find(qn("w:rPr"))
        if properties is None:
            return style
        changes: dict[str, Any] = {}
        for flag, name in (("bold", "w:b"), ("italic", "w:i"), ("strike", "w:strike")):
            value = _docx_toggle(properties, qn(name), qn("w:val"))
            if value is not None:
                changes[flag] = value
        if _docx_toggle(properties, qn("w:dstrike"), qn("w:val")):
            changes["strike"] = True
        underline = properties.find(qn("w:u"))
        if underline is not None:
            changes["underline"] = (underline.get(qn("w:val")) or "single") != "none"
        vertical = properties.find(qn("w:vertAlign"))
        if vertical is not None:
            where = vertical.get(qn("w:val")) or ""
            changes["superscript"] = where == "superscript"
            changes["subscript"] = where == "subscript"
        fonts = properties.find(qn("w:rFonts"))
        if fonts is not None:
            name = _font_name(fonts.get(qn("w:ascii")) or fonts.get(qn("w:hAnsi")))
            if name and name.lower() in MONOSPACE_FONTS:
                changes["monospace"] = True
                changes["font"] = None
            elif name:
                changes["monospace"] = False
                changes["font"] = name
        size = properties.find(qn("w:sz"))
        if size is not None:
            half = _number(size.get(qn("w:val")))
            if half:
                changes["size"] = round(half / 2.0, 2)
        color = properties.find(qn("w:color"))
        if color is not None:
            changes["color"] = css.color("#" + (color.get(qn("w:val")) or ""))
        highlight = properties.find(qn("w:highlight"))
        if highlight is not None:
            changes["background"] = css.color(highlight.get(qn("w:val")) or "")
        shading = properties.find(qn("w:shd"))
        if shading is not None:
            fill = css.color("#" + (shading.get(qn("w:fill")) or ""))
            if fill:
                changes["background"] = fill
        return replace(style, **changes) if changes else style

    def _drop(self, what: str) -> None:
        """Count one thing the document has no place for."""
        self._dropped[what] = self._dropped.get(what, 0) + 1


def _docx_style_name(paragraph) -> str:
    """The name of the style one paragraph is written in, or an empty string."""
    try:
        return str(getattr(paragraph.style, "name", "") or "")
    except (AttributeError, KeyError, ValueError) as error:
        logger.debug("a paragraph names a style the package does not declare (%s)", error)
        return ""


def _docx_heading(name: str) -> int:
    """The heading level one Word style name states, or 0."""
    text = str(name or "").strip().lower()
    if text == "title":
        return 1
    if text == "subtitle":
        return 2
    match = _HEADING_STYLE.match(text)
    return int(match.group(1)) if match else 0


def _docx_list_style(name: str) -> tuple[int, bool]:
    """The list level and marker one Word style name states, or ``(0, False)``."""
    match = _LIST_STYLE.match(str(name or "").strip().lower())
    if not match:
        return 0, False
    level = int(match.group(2) or 1)
    return max(1, min(level, MAX_STYLE_DEPTH)), match.group(1) == "number"


def _docx_numbering(element, qn) -> tuple[int, int] | None:
    """The numbering definition and level one paragraph names, or ``None``."""
    properties = element.find(qn("w:pPr"))
    if properties is None:
        return None
    numbering = properties.find(qn("w:numPr"))
    if numbering is None:
        return None
    identifier = numbering.find(qn("w:numId"))
    number = _number(identifier.get(qn("w:val"))) if identifier is not None else None
    if not number:
        return None
    level = numbering.find(qn("w:ilvl"))
    depth = _number(level.get(qn("w:val"))) if level is not None else 0.0
    return int(number), max(0, min(int(depth or 0), MAX_STYLE_DEPTH))


def _numbering_formats(word, qn) -> dict[tuple[int, int], str]:
    """How each numbered list in a ``.docx`` marks each of its levels."""
    try:
        root = word.part.numbering_part.element
    except (AttributeError, KeyError, NotImplementedError, ValueError):
        return {}
    abstract: dict[tuple[str, str], str] = {}
    for definition in root.findall(qn("w:abstractNum")):
        identifier = definition.get(qn("w:abstractNumId"))
        for level in definition.findall(qn("w:lvl")):
            marker = level.find(qn("w:numFmt"))
            if marker is not None:
                abstract[(identifier, level.get(qn("w:ilvl")))] = marker.get(qn("w:val")) or ""
    found: dict[tuple[int, int], str] = {}
    for used in root.findall(qn("w:num")):
        reference = used.find(qn("w:abstractNumId"))
        number = _number(used.get(qn("w:numId")))
        if reference is None or number is None:
            continue
        wanted = reference.get(qn("w:val"))
        for (identifier, level), marker in abstract.items():
            depth = _number(level)
            if identifier == wanted and depth is not None:
                found[(int(number), int(depth))] = marker
    return found


def _docx_toggle(properties, tag: str, attribute: str) -> bool | None:
    """Whether one on-or-off run property is set, or ``None`` where it is not declared."""
    element = properties.find(tag)
    if element is None:
        return None
    value = element.get(attribute)
    if value is None:
        return True
    return str(value).lower() not in ("0", "false", "off")


def _docx_align(element, qn) -> str | None:
    """The alignment one paragraph declares, as CSS spells it."""
    properties = element.find(qn("w:pPr"))
    if properties is None:
        return None
    justification = properties.find(qn("w:jc"))
    if justification is None:
        return None
    value = str(justification.get(qn("w:val")) or "").strip().lower()
    return css.alignment(DOCX_ALIGNMENTS.get(value, value))


def _docx_indent(element, qn) -> float | None:
    """How far one paragraph is indented from the left, in points."""
    properties = element.find(qn("w:pPr"))
    if properties is None:
        return None
    indent = properties.find(qn("w:ind"))
    if indent is None:
        return None
    twips = _number(indent.get(qn("w:left")) or indent.get(qn("w:start")))
    return round(twips / TWIPS_PER_POINT, 2) if twips and twips > 0 else None


def _docx_rule(element, qn) -> bool:
    """Whether one empty paragraph is drawn as a line across the page."""
    properties = element.find(qn("w:pPr"))
    if properties is None:
        return False
    borders = properties.find(qn("w:pBdr"))
    bottom = borders.find(qn("w:bottom")) if borders is not None else None
    if bottom is None:
        return False
    return str(bottom.get(qn("w:val")) or "single").lower() not in ("none", "nil")


def _docx_cell_spans(cell, qn) -> tuple[int, Any]:
    """How many columns one cell covers, and the vertical merge it takes part in."""
    properties = cell.find(qn("w:tcPr"))
    if properties is None:
        return 1, None
    span = properties.find(qn("w:gridSpan"))
    columns = _number(span.get(qn("w:val"))) if span is not None else None
    return max(1, int(columns or 1)), properties.find(qn("w:vMerge"))


def _docx_drawing(element, qn) -> tuple[float | None, float | None, str]:
    """The drawn size in points and the description one drawing carries."""
    width = height = None
    extent = element.find(".//" + qn("wp:extent"))
    if extent is not None:
        width = _points(extent.get("cx"))
        height = _points(extent.get("cy"))
    alt = ""
    properties = element.find(".//" + qn("wp:docPr"))
    if properties is not None:
        alt = str(properties.get("descr") or properties.get("title") or "").strip()
    return width, height, alt


def _docx_metadata(word, payload: bytes) -> Metadata:
    """The metadata a ``.docx`` package carries, core and custom properties together."""
    properties = word.core_properties
    custom, application = _docx_properties(payload)
    rights = _taken(custom, "copyright")
    generator = _taken(custom, "generator") or application
    for field, value in (
        ("Subject", getattr(properties, "subject", "")),
        ("Category", getattr(properties, "category", "")),
    ):
        text = str(value or "").strip()
        if text:
            custom.setdefault(field, text)
    return Metadata(
        title=_docx_text(properties, "title"),
        description=_docx_text(properties, "comments"),
        author=_docx_text(properties, "author"),
        copyright=rights,
        language=_docx_text(properties, "language"),
        keywords=_docx_text(properties, "keywords"),
        created=_stamp(getattr(properties, "created", None)),
        modified=_stamp(getattr(properties, "modified", None)),
        generator=generator,
        custom=custom,
    )


def _docx_text(properties, field: str) -> str:
    """One core property as text, or an empty string."""
    try:
        return str(getattr(properties, field, "") or "").strip()
    except (AttributeError, ValueError) as error:
        logger.debug("the %s property could not be read (%s)", field, error)
        return ""


def _docx_properties(payload: bytes) -> tuple[dict[str, str], str]:
    """The custom properties and the program named inside a ``.docx`` package."""
    found: dict[str, str] = {}
    application = ""
    try:
        with zipfile.ZipFile(BytesIO(payload)) as archive:
            names = set(archive.namelist())
            if ooxml.CUSTOM_PART in names:
                root = _parsed(archive.read(ooxml.CUSTOM_PART), ooxml.CUSTOM_PART)
                for element in root if root is not None else ():
                    if len(found) >= ooxml.MAX_PROPERTIES:
                        break
                    name = str(element.get("name") or "").strip()
                    value = "".join(child.text or "" for child in element).strip()
                    if name and value:
                        found.setdefault(name, value)
            if APP_PART in names:
                root = _parsed(archive.read(APP_PART), APP_PART)
                for element in root if root is not None else ():
                    if element.tag.rpartition("}")[2] == "Application":
                        application = str(element.text or "").strip()
                        break
    except (KeyError, OSError, ValueError, zipfile.BadZipFile) as error:
        logger.debug("the properties of this %s file were not read (%s)", DOCX, error)
    return found, application


def _taken(custom: dict[str, str], name: str) -> str:
    """Pull one named pair out of the custom properties, whatever its case."""
    for key in list(custom):
        if key.strip().lower() == name:
            return custom.pop(key)
    return ""


def _number(value: Any) -> float | None:
    """One attribute as a number, or ``None`` where it is not one."""
    try:
        return float(str(value).strip())
    except (TypeError, ValueError):
        return None


def _points(value: Any) -> float | None:
    """One length in English Metric Units as points, or ``None``."""
    number = _number(value)
    return round(number / EMU_PER_POINT, 2) if number and number > 0 else None


# -------------------------------------------------------------------------- .odt


def _read_odt(payload: bytes) -> container.Document:
    """Read an ``.odt`` package through odfdo.

    Args:
        payload: The package bytes.

    Returns:
        The document.

    Raises:
        NotADocument: The package is another kind of OpenDocument file, or odfdo would not
            open it.
        DependencyError: odfdo is missing or unusable.
    """
    deps.require("odfdo", feature=FEATURE)
    import odfdo

    _odt_kind(payload)
    try:
        text = odfdo.Document(BytesIO(payload))
        body = text.body
    except Exception as error:
        # odfdo raises its own exception types for a package it cannot open.
        raise container.NotADocument(
            f"this {ODT} file could not be opened by odfdo ({error}).\n"
            f"  The package is damaged, or one of its parts is not the XML the format "
            f"describes. Open it in a word processor and save it again as {ODT}."
        ) from error
    if body is None:
        raise container.NotADocument(
            f"this {ODT} file holds no document body, so there is nothing in it to read."
        )
    reader = _Odt(text)
    reader.body()
    reader.assets.report(ODT)
    reader.report()
    return container.Document(
        _cleaned(reader.page.result(), ODT), _odt_metadata(text), reader.assets.found
    )


def _odt_kind(payload: bytes) -> None:
    """Confirm the package is a text document rather than another OpenDocument kind.

    Args:
        payload: The package bytes.

    Raises:
        NotADocument: It is a spreadsheet, a presentation or a drawing.
    """
    try:
        with zipfile.ZipFile(BytesIO(payload)) as archive:
            declared = archive.read(container.MIMETYPE_ENTRY).decode("ascii", "replace")
    except (KeyError, OSError, ValueError, zipfile.BadZipFile):
        return
    declared = declared.strip()
    if not declared.startswith(ODT_MIMETYPE):
        raise container.NotADocument(
            f"this file says its type is {declared}, which is an OpenDocument file but "
            f"not a text document.\n"
            f"  A document carries writing. A spreadsheet, a presentation and a drawing "
            f"each hold something else, and this node opens none of them. Open it in its "
            f"own program and save the part wanted as {ODT}."
        )


class _Odt:
    """Reads one odfdo document into markup and pictures.

    Attributes:
        text: The document being read.
        page: The markup being assembled.
        assets: The pictures found so far.
    """

    def __init__(self, text) -> None:
        """Read one odfdo document.

        Args:
            text: The opened ``Document``.
        """
        self.text = text
        self.page = _Markup()
        self.assets = _Assets()
        self._styles: dict[tuple[str, str], tuple[tuple[str, ...], dict[str, str]]] = {}
        self._lists: dict[str, bool] = {}
        self._dropped: dict[str, int] = {}

    def body(self) -> None:
        """Read every block of the document body."""
        for child in self.text.body.children:
            self.block(child)

    def report(self) -> None:
        """Log whatever the file holds that a document has no place for."""
        for what, count in sorted(self._dropped.items()):
            logger.info("%d %s in this %s file were not read.", count, what, ODT)

    def block(self, element) -> None:
        """Read one element of a body, a section or a table cell."""
        tag = element.tag
        if tag == "text:h":
            self.paragraph(element, heading=True)
        elif tag == "text:p":
            self.paragraph(element)
        elif tag == "text:list":
            self.listing(element, 1)
        elif tag == "table:table":
            self.table(element)
        elif tag in ODT_GROUPS:
            for child in element.children:
                self.block(child)
        elif tag not in ODT_IGNORED:
            self._drop(f"{tag} element(s)")

    def paragraph(self, element, heading: bool = False) -> None:
        """Read one paragraph as whatever kind of block its style makes it."""
        names, properties = self.resolve("paragraph", _odt_style(element))
        style = _block_style(
            css.alignment(properties.get("fo:text-align", "")),
            css.length(properties.get("fo:margin-left", "")),
        )
        base = _odt_format(properties, Style())
        if _odt_is(names, ODF_PREFORMATTED):
            self.page.preformatted(self.written(element, replace(base, monospace=True)), style)
            return
        inner = self.written(element, base)
        if heading:
            level = _number(element.get_attribute("text:outline-level"))
            self.page.heading(int(level or _odt_heading(names) or 1), inner, style)
            return
        if not inner.strip() and _odt_rule(properties):
            self.page.rule(style)
            return
        if _odt_is(names, ODF_CAPTION):
            self.page.caption(inner)
            return
        self.page.paragraph(inner, style, quote=_odt_is(names, ODF_QUOTE))

    def written(self, element, style: Style) -> str:
        """The markup of everything written inside one paragraph or heading."""
        runs = _Runs()
        self.content(element, style, runs)
        return runs.result()

    def content(self, element, style: Style, runs: _Runs) -> None:
        """Read everything written inside one paragraph, heading or span."""
        runs.text(element.text or "", style)
        for child in element.children:
            self.inline(child, style, runs)
            runs.text(child.tail or "", style)

    def inline(self, element, style: Style, runs: _Runs) -> None:
        """Read one element written inside a paragraph."""
        tag = element.tag
        if tag == "text:span":
            name = _odt_style(element)
            inner = _odt_format(self.resolve("text", name)[1], style) if name else style
            self.content(element, inner, runs)
        elif tag == "text:a":
            target = str(element.get_attribute("xlink:href") or "").strip()
            self.content(element, replace(style, link=target or None), runs)
        elif tag == "text:s":
            count = _number(element.get_attribute("text:c")) or 1
            runs.text(" " * max(1, min(int(count), MAX_SPACES)), style)
        elif tag == "text:tab":
            runs.text("\t", style)
        elif tag == "text:line-break":
            runs.element("<br>")
        elif tag in ("draw:frame", "draw:image"):
            runs.element(self.picture(element, style))
        elif tag in ODT_DROPPED_INLINE:
            self._drop(ODT_DROPPED_INLINE[tag])
        elif tag not in ODT_SKIPPED_INLINE:
            self.content(element, style, runs)

    def listing(self, element, level: int, marker: bool | None = None) -> None:
        """Read one list and every list nested inside it.

        Args:
            element: The ``text:list`` to read.
            level: How deep the list sits, the outermost being 1.
            marker: Whether the list around this one is numbered, taken where this list
                names no style of its own.
        """
        ordered = self._ordered(element, marker)
        for item in element.children:
            if item.tag in ("text:list-item", "text:list-header"):
                self._item(item, level, ordered)

    def table(self, element) -> None:
        """Read one table, its spans rebuilt and its covered cells left out."""
        rows: list[str] = []
        self._rows(element, rows, False)
        if any(row != "<tr></tr>" for row in rows):
            self.page.table("".join(rows))

    def cell(self, element) -> str:
        """Read one cell into markup of its own."""
        outer = self.page
        self.page = _Markup()
        try:
            for child in element.children:
                self.block(child)
            return self.page.result()
        finally:
            self.page = outer

    def picture(self, element, style: Style) -> str:
        """Read one picture, drawn in the line where the package carries the file."""
        width = css.length(str(element.get_attribute("svg:width") or ""))
        height = css.length(str(element.get_attribute("svg:height") or ""))
        image = element if element.tag == "draw:image" else None
        alt = ""
        if image is None:
            for child in element.children:
                if child.tag == "draw:image" and image is None:
                    image = child
                elif child.tag in ("svg:desc", "svg:title") and not alt:
                    alt = str(child.text or "").strip()
        source = str(image.get_attribute("xlink:href") or "") if image is not None else ""
        if source and _remote(source):
            return _picture_markup(source, alt, width, height)
        data = self._part(source) if source else _odt_embedded(image)
        name = self.assets.add(source or f"picture{len(self.assets.found) + 1}", data, source)
        if name is None:
            self._drop("picture(s) the package does not carry")
            return _inline(f"[{alt or 'picture'}]", replace(style, italic=True))
        return _picture_markup(container.ASSET_PREFIX + name, alt, width, height)

    def resolve(self, family: str, name: str) -> tuple[tuple[str, ...], dict[str, str]]:
        """Every style one name inherits from, and the properties they declare together.

        Args:
            family: ``"text"``, ``"paragraph"`` or ``"list"``.
            name: The style named on an element.

        Returns:
            ``(names nearest first, properties)``. A name in :data:`ODF_BASE_STYLES`
            contributes its place in the chain and none of its properties.
        """
        key = (family, name)
        if key in self._styles:
            return self._styles[key]
        names: list[str] = []
        styles: list[Any] = []
        current = str(name or "")
        seen: set[str] = set()
        while current and current not in seen and len(seen) < MAX_STYLE_DEPTH:
            seen.add(current)
            readable = _odt_name(current)
            names.append(readable)
            style = self._style(family, current)
            if style is None:
                break
            if not _odt_base(readable):
                styles.append(style)
            current = str(getattr(style, "parent_style", "") or "")
        found: dict[str, str] = {}
        for style in reversed(styles):
            for area in ("paragraph", "text", "graphic"):
                found.update(_odt_area(style, area))
        self._styles[key] = (tuple(names), found)
        return self._styles[key]

    def _item(self, item, level: int, ordered: bool) -> None:
        """Read one list item, its paragraphs first and its nested lists inside it."""
        pieces: list[str] = []
        written = False
        nested = False
        for child in item.children:
            if child.tag == "text:list":
                if pieces:
                    self.page.item(level, ordered, "<br>".join(pieces))
                    pieces = []
                    written = True
                nested = True
                self.listing(child, level + 1, ordered)
            elif child.tag in ("text:p", "text:h"):
                properties = self.resolve("paragraph", _odt_style(child))[1]
                pieces.append(self.written(child, _odt_format(properties, Style())))
            else:
                self.block(child)
        if pieces or not (written or nested):
            self.page.item(level, ordered, "<br>".join(pieces))

    def _rows(self, element, rows: list[str], header: bool) -> None:
        """Collect every row of a table, reading through whatever groups them."""
        for child in element.children:
            tag = child.tag
            if tag == "table:table-row":
                rows.extend(self._row(child, header))
            elif tag == "table:table-header-rows":
                self._rows(child, rows, True)
            elif tag in ("table:table-row-group", "table:table-rows"):
                self._rows(child, rows, header)

    def _row(self, row, header: bool) -> list[str]:
        """One table row, written out as many times as it repeats."""
        cells: list[str] = []
        for cell in row.children:
            if cell.tag != "table:table-cell":
                continue
            colspan = _odt_span(cell, "table:number-columns-spanned")
            rowspan = _odt_span(cell, "table:number-rows-spanned")
            html = self.cell(cell)
            repeat = _odt_repeat(cell, "table:number-columns-repeated")
            cells.extend([_cell(html, header, colspan, rowspan)] * repeat)
        # A row every span reaches into keeps its place, empty.
        written = "<tr>" + "".join(cells) + "</tr>"
        return [written] * _odt_repeat(row, "table:number-rows-repeated")

    def _ordered(self, element, marker: bool | None = None) -> bool:
        """Whether one list is numbered rather than bulleted.

        Args:
            element: The ``text:list`` being read.
            marker: What the list around this one is marked with.

        Returns:
            True where the items are numbered. A list naming no style of its own
            carries the marker of the list around it.
        """
        name = _odt_style(element)
        if not name:
            return bool(marker)
        if name in self._lists:
            return self._lists[name]
        ordered = "number" in name.lower()
        style = self._style("list", name)
        if style is not None:
            for child in style.children:
                if child.tag.endswith("-style-number"):
                    ordered = True
                    break
                if child.tag.endswith("-style-bullet"):
                    ordered = False
                    break
        self._lists[name] = ordered
        return ordered

    def _style(self, family: str, name: str):
        """One named style out of the document, or ``None``."""
        try:
            return self.text.get_style(family, name)
        except Exception as error:
            logger.debug("the %s style %r could not be read (%s)", family, name, error)
            return None

    def _part(self, source: str) -> bytes:
        """One picture out of the package, or no bytes at all."""
        try:
            data = self.text.get_part(source)
        except Exception as error:
            logger.debug("a picture is not in the package (%s)", error)
            return b""
        return data if isinstance(data, bytes) else b""

    def _drop(self, what: str) -> None:
        """Count one thing the document has no place for."""
        self._dropped[what] = self._dropped.get(what, 0) + 1


def _odt_style(element) -> str:
    """The style name one element carries, or an empty string."""
    return str(element.get_attribute("text:style-name") or "")


def _odt_name(name: str) -> str:
    """One ODF style name with its escaped spaces put back."""
    return str(name or "").replace(ODF_SPACE, " ").strip()


def _odt_base(name: str) -> bool:
    """Whether a style is one of the document defaults, whose properties are not read."""
    lowered = name.lower()
    return lowered in ODF_BASE_STYLES or bool(_HEADING_STYLE.match(lowered))


def _odt_is(names: tuple[str, ...], wanted: tuple[str, ...]) -> bool:
    """Whether a style, or anything it inherits from, is one of the named kinds."""
    return any(name.lower() in wanted for name in names)


def _odt_heading(names: tuple[str, ...]) -> int:
    """The heading level a style chain states, or 0."""
    for name in names:
        match = _HEADING_STYLE.match(name.lower())
        if match:
            return int(match.group(1))
    return 0


def _odt_rule(properties: Mapping[str, str]) -> bool:
    """Whether one empty paragraph is drawn as a line across the page."""
    return any(
        key.startswith("fo:border") and value and value != "none"
        for key, value in properties.items()
    )


def _odt_span(cell, attribute: str) -> int:
    """How far one cell spans, at least 1."""
    return max(1, min(int(_number(cell.get_attribute(attribute)) or 1), MAX_SPAN))


def _odt_repeat(element, attribute: str) -> int:
    """How many times one row or cell is written out, at least 1."""
    return max(1, min(int(_number(element.get_attribute(attribute)) or 1), MAX_REPEAT))


def _odt_area(style, area: str) -> dict[str, str]:
    """One area of properties off an ODF style, or nothing at all."""
    try:
        found = style.get_properties(area)
    except Exception as error:
        logger.debug("the %s properties of a style could not be read (%s)", area, error)
        return {}
    return {str(key): str(value) for key, value in (found or {}).items()}


def _odt_format(properties: Mapping[str, str], style: Style) -> Style:
    """The character formatting ODF properties declare, over the formatting around them."""
    changes: dict[str, Any] = {}
    weight = properties.get("fo:font-weight", "")
    if weight:
        changes["bold"] = weight not in ("normal", "100", "200", "300", "400")
    posture = properties.get("fo:font-style", "")
    if posture:
        changes["italic"] = posture != "normal"
    underline = properties.get("style:text-underline-style", "")
    if underline:
        changes["underline"] = underline != "none"
    struck = properties.get("style:text-line-through-style", "")
    if struck:
        changes["strike"] = struck != "none"
    position = properties.get("style:text-position", "").strip().lower()
    if position:
        first = position.split()[0]
        raised = _number(first.rstrip("%"))
        changes["superscript"] = first.startswith("super") or bool(raised and raised > 0)
        changes["subscript"] = first.startswith("sub") or bool(raised and raised < 0)
    color = css.color(properties.get("fo:color", ""))
    if color:
        changes["color"] = color
    background = css.color(properties.get("fo:background-color", ""))
    if background:
        changes["background"] = background
    name = _font_name(
        properties.get("style:font-name") or properties.get("fo:font-family")
    )
    if name and name.lower() in MONOSPACE_FONTS:
        changes["monospace"] = True
        changes["font"] = None
    elif name:
        changes["monospace"] = False
        changes["font"] = name
    size = css.length(properties.get("fo:font-size", ""))
    if size:
        changes["size"] = round(size, 2)
    return replace(style, **changes) if changes else style


def _odt_embedded(image) -> bytes:
    """The bytes one picture carries inside the markup rather than as a package part."""
    if image is None:
        return b""
    for child in image.children:
        if child.tag == "office:binary-data":
            try:
                return base64.b64decode("".join((child.text or "").split()), validate=False)
            except (binascii.Error, ValueError) as error:
                logger.debug("a picture written into the markup could not be read (%s)", error)
    return b""


def _odt_metadata(text) -> Metadata:
    """The metadata an ``.odt`` package carries, its user-defined pairs included."""
    meta = getattr(text, "meta", None)
    if meta is None:
        logger.debug("this %s file carries no readable meta.xml", ODT)
        return Metadata()
    pairs = _odt_field(meta, "user_defined_metadata") or {}
    custom = {str(name): str(value) for name, value in pairs.items() if value}
    rights = _taken(custom, "copyright")
    subject = _odt_text(meta, "subject")
    if subject:
        custom.setdefault("Subject", subject)
    return Metadata(
        title=_odt_text(meta, "title"),
        description=_odt_text(meta, "description"),
        author=_odt_text(meta, "creator"),
        copyright=rights,
        language=_odt_text(meta, "language"),
        keywords=_odt_text(meta, "keywords"),
        created=_stamp(_odt_field(meta, "creation_date")),
        modified=_stamp(_odt_field(meta, "modification_date")),
        generator=_odt_text(meta, "generator"),
        custom=custom,
    )


def _odt_field(meta, field: str) -> Any:
    """One metadata field, by whichever spelling the installed odfdo offers."""
    declared = getattr(type(meta), field, None)
    if isinstance(declared, property):
        try:
            return getattr(meta, field)
        except Exception as error:
            logger.debug("meta.%s could not be read (%s)", field, error)
    getter = getattr(meta, f"get_{field}", None)
    if callable(getter):
        try:
            return getter()
        except Exception as error:
            logger.debug("meta.get_%s could not be read (%s)", field, error)
    return None


def _odt_text(meta, field: str) -> str:
    """One metadata field as text."""
    value = _odt_field(meta, field)
    if isinstance(value, (list, tuple, set)):
        return ", ".join(str(entry) for entry in value if entry)
    return str(value or "").strip()
