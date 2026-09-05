"""Read one document file from disk onto a DOC socket."""

from __future__ import annotations

from pathlib import Path

from comfy_api.latest import io

from ...modules import config, deps
from ...modules.archive import kinds, picks
from ...modules.compat.types import DOC
from ...modules.document import container
from ...modules.log import get_logger
from ...modules.util import file_listing, sandbox

logger = get_logger("nodes.document")

#: Menu entry shown when none of the three folders holds a document this node reads.
NO_DOCUMENTS = "No Documents"

#: How many entries the menu offers.
MAX_OPTIONS = 500


def extensions() -> tuple[str, ...]:
    """Every suffix the menu lists, the container's own and the two it converts."""
    from ...modules.document import office

    return (container.SUFFIX, *office.FORMATS)


def options() -> list[str]:
    """The menu's entries, or ``[NO_DOCUMENTS]`` when there are none."""
    return list(
        file_listing.labels(extensions(), file_listing.TAGS, MAX_OPTIONS)
    ) or [NO_DOCUMENTS]


class LoadDocument(io.ComfyNode):
    """One document, read from the path a widget names."""

    @classmethod
    def define_schema(cls) -> io.Schema:
        return io.Schema(
            node_id="WASLoadDocument",
            display_name="Load Document",
            search_aliases=[
                "WASLoadDocument",
                "Load Document",
                "load doc",
                "open document",
                "read document",
                "load wasdoc",
                "wasdoc",
                "load docx",
                "open word document",
                "read word file",
                "import docx",
                "load odt",
                "open odt",
                "opendocument",
            ],
            category="WAS Suite/Document",
            description=(
                (
                    "Open one document from disk and put it on a DOC wire, with its markup, "
                    "its metadata and every file embedded in it. The file menu lists every "
                    "document in ComfyUI's input, output and temp folders, three folders "
                    "deep, and any folder added under paths.allow_read. A "
                    ".wasdoc is the container Save DOC writes and arrives exactly as the file "
                    "holds it. A .docx and an .odt open too, converted rather than copied: "
                    "the writing, tables, links and pictures come across, page geometry, "
                    "headers, footnotes, comments and tracked changes do not. Those two read "
                    "through python-docx and odfdo, named in the error when either is "
                    "missing, and features.document_export can refuse them. A .pdf is never "
                    "read. The document is read again whenever the file changes."
                )
            ),
            inputs=[
                io.Combo.Input(
                    "file",
                    options=options(),
                    tooltip=(
                        "Which document to open. The menu lists every .wasdoc, .docx and "
                        ".odt in ComfyUI's input, output and temp folders, tagged '[input]', "
                        "'[output]' or '[temp]'. A folder added under paths.allow_read in "
                        "config.yaml appears under its own name."
                    ),
                ),
            ],
            outputs=[
                DOC.Output(
                    tooltip=(
                        "The whole document on one wire: its markup, its metadata and its "
                        "embedded files. A .wasdoc arrives exactly as the file holds it; a "
                        ".docx or .odt arrives converted, its writing and pictures kept and "
                        "its page layout dropped. Wire it into View DOC Metadata, Convert "
                        "DOC to Plaintext, or Save DOC."
                    ),
                ),
            ],
        )

    @classmethod
    def fingerprint_inputs(cls, file=""):
        """Read again when the file on disk has changed."""
        return picks.fingerprint(_named(file))

    @classmethod
    def execute(cls, file="") -> io.NodeOutput:
        """Read the file and hand on the document it holds.

        Raises:
            DocumentError: ``path`` is empty, names a folder, names nothing that is there,
                or names a file that unpacks past the bounds a document is read under.
            NotADocument: It names a zip archive, a file of a kind this pack does not read
                as a document, or a file whose bytes are not that format at all.
            UnsupportedVersion: The container names a layout this build does not read.
            ValueError: The format is read through a document library and
                ``features.document_export`` is false.
            DependencyError: That library is missing or unusable.
            PathNotAllowed: It resolved outside every permitted read root.
            OSError: The file is there and could not be read.
        """
        from ...modules.document import office

        resolved = _resolved(_named(file))
        extension = resolved.suffix.lower()
        converted = office.reads(extension)
        if converted:
            _require_export_group(extension)
        try:
            data = resolved.read_bytes()
            document = (
                office.read(data, extension)
                if converted
                else container.Document.from_bytes(data)
            )
        except container.DocumentError as error:
            raise type(error)(f"{resolved} could not be read.\n  {error}") from error

        _report(resolved, document, extension if converted else "")
        return io.NodeOutput(document)


def _named(file: str) -> str:
    """The document the menu names.

    Args:
        file: The menu label, such as ``report.wasdoc [input]``.

    Returns:
        The path the label resolves to, the label itself where it does not resolve, and an
        empty string where nothing was chosen.
    """
    label = str(file or "").strip()
    if not label or label == NO_DOCUMENTS:
        return ""
    return file_listing.resolve(label, extensions(), file_listing.TAGS) or label


def _require_export_group(extension: str) -> None:
    """Refuse a format read through a document library while its group is set to false.

    Args:
        extension: One of :data:`modules.document.office.FORMATS`.

    Raises:
        ValueError: ``features.document_export`` is false, naming the setting, the install
            command and the format that needs neither.
    """
    from ...modules.document import office

    if config.load_config()["features"]["document_export"]:
        return
    packages = [deps.PIP_NAMES.get(name, name) for name in office.PACKAGES.values()]
    named = deps.PIP_NAMES.get(office.PACKAGES[extension], office.PACKAGES[extension])
    raise ValueError(
        f"a {extension} file is read through {named}, and the {office.FEATURE} group is "
        f"set to false in your config, so this pack will not open one.\n"
        f"  That group ships off, so this is the one thing to turn on to read these\n"
        f"  formats. Put this in config.yaml and restart ComfyUI:\n"
        f"      features:\n"
        f"        document_export: true\n"
        f"  Then install the two libraries, both permissive-licensed and pure python:\n"
        f"      {deps.install_command(*packages, feature=office.FEATURE)}\n"
        f"  Or open a {container.SUFFIX}, which needs nothing installed and keeps "
        f"everything a document holds."
    )


def _report(resolved: Path, document: container.Document, converted: str) -> None:
    """Log what was read, and what a conversion left behind.

    Args:
        resolved: The file that was read.
        document: What it turned into.
        converted: The format it was converted from, or an empty string for a container.
    """
    if not document.plain_text:
        logger.warning(
            "Load Document opened %s and it holds nothing a reader would see, so the DOC "
            "output carries its metadata and no text. The file was read; its content is "
            "%d character(s) of markup.",
            resolved, len(document.content),
        )
        return
    logger.info(
        "Load Document read %s: %d word(s), %d embedded file(s)",
        resolved, document.word_count, len(document.assets),
    )
    if converted:
        logger.info(
            "%s was converted into a document rather than copied. Its writing, headings, "
            "lists, tables, links, pictures and metadata came across. The page size and "
            "margins, headers and footers, footnotes and endnotes, comments, tracked "
            "changes, text boxes, charts and any styling a run does not carry for itself "
            "did not: a document is HTML and has nowhere to keep them. Saving it back to "
            "%s lays it out afresh on the page Save DOC is set to.",
            resolved, converted,
        )


def _resolved(value: str) -> Path:
    """The document one widget value names, resolved and confirmed to be a file there.

    Args:
        value: The raw widget value.

    Returns:
        The absolute path, inside a permitted read root.

    Raises:
        DocumentError: The value is empty, names a folder, or names nothing that is there.
        NotADocument: It names a zip archive, or a file of a kind this pack does not read
            as a document.
        PathNotAllowed: It resolved outside every permitted read root.
    """
    from ...modules.document import office

    text = str(value).strip()
    if not text:
        raise container.DocumentError(
            f"no document was given.\n"
            f"  Pick one in the file menu, which lists every document in ComfyUI's input, "
            f"output and temp folders, or in a folder added under paths.allow_read in "
            f"config.yaml. The extensions read are {_extensions()}."
        )
    path = sandbox.resolve_read(text)
    if path.is_dir():
        raise container.DocumentError(
            f"{path} is a folder rather than a document.\n"
            f"  This node opens one file. To read every document inside a zip archive, use "
            f"Load Documents from ZIP instead."
        )
    if not path.exists():
        raise container.DocumentError(
            f"the document {path} cannot be found.\n"
            f"  A path with no folders in it is read against the folder ComfyUI was started "
            f"in and nowhere else, so give the whole path, or put the file in ComfyUI's "
            f"input folder and name it 'input/{path.name}'."
        )
    if path.suffix.lower() == ".zip":
        raise container.NotADocument(
            f"{path} is a zip archive rather than one document.\n"
            f"  To read every document inside it, use Load Documents from ZIP instead."
        )
    if not office.reads(path.suffix) and kinds.kind_of(path.name) != kinds.DOCUMENT:
        raise container.NotADocument(
            f"{path} is not a document this pack reads. The extensions it reads are "
            f"{_extensions()}: the container Save DOC writes, and the two word processor "
            f"formats it also writes.\n"
            f"  A .pdf cannot be read back into a document, and neither can the older "
            f"binary .doc. An HTML page or a text file becomes one by going through Load "
            f"Text File and then Text to DOC."
        )
    return path


def _extensions() -> str:
    """Every extension this node opens, in one comma-separated string."""
    from ...modules.document import office

    return ", ".join((container.SUFFIX, *office.FORMATS))
